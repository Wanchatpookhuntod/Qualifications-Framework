"""Word (.docx) exporters for TQF3 (มคอ.3) and TQF5 (มคอ.5) documents.

These builders mirror the read-only layouts in
``templates/shared/tqf3_readonly.html`` and ``templates/shared/tqf5_readonly.html``
so the downloaded Word file matches what reviewers see on screen. Keep the field
keys and fallback order in sync with those templates.

The functions return an in-memory ``io.BytesIO`` ready to hand to Flask's
``send_file``. No Firestore access happens here; callers resolve the section /
course / term context and pass plain dictionaries in.
"""

from __future__ import annotations

import base64
import io
import re
from datetime import date
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Mm, Pt, RGBColor
from docx.table import _Cell

# Font that ships with most Thai Office installs and renders Thai cleanly.
THAI_FONT = "TH Sarabun New"
_BLACK = RGBColor(0x00, 0x00, 0x00)
# Headings are black; muted notes/subtitles keep their light gray.
_HEADING_COLOR = _BLACK
_MUTED_COLOR = RGBColor(0x55, 0x55, 0x55)
_DASH = "-"

# Printable width on A4 (210mm) with 0.5" left + 0.5" right margins. A small
# slack is kept so rounding can't push a fixed table past the right margin.
_CONTENT_WIDTH_IN = (210 / 25.4) - 1.0 - 0.02

# Thai word-wrap support (see format_tqf_new/word_export_format_thai/). Word breaks
# Thai lines at word boundaries only when (1) runs are tagged th-TH and (2) a
# zero-width space (U+200B) is inserted between words. Without (2), narrow table
# cells split Thai mid-word / spread it character-by-character.
_THAI_RE = re.compile(r"[฀-๿]")
_ZWSP = "​"

# Characters XML 1.0 (and therefore .docx) forbids: NULL and C0/C1 control codes
# other than tab/newline/carriage-return. Pasting from PDFs/Word often smuggles
# these in (e.g. \x0b, \x0c, \x00); python-docx then raises ValueError when it
# writes the run, 500-ing the export. Strip them before any text reaches a run.
_XML_ILLEGAL_RE = re.compile(
    "[^\u0009\u000a\u000d\u0020-\ud7ff\ue000-\ufffd\U00010000-\U0010ffff]"
)


def _xml_safe(text: str) -> str:
    """Drop characters that aren't legal in XML/.docx so a run can't crash."""
    return _XML_ILLEGAL_RE.sub("", text)


def _segment_thai(text: str) -> str:
    """Insert U+200B between Thai words so Word can wrap at word boundaries."""
    if not text or not _THAI_RE.search(text):
        return text
    try:
        from pythainlp.tokenize import word_tokenize
    except Exception:  # pragma: no cover - pythainlp optional at runtime
        return text
    try:
        tokens = word_tokenize(text, keep_whitespace=True)
    except Exception:  # pragma: no cover - defensive
        return text
    out: List[str] = []
    for tok in tokens:
        out.append(tok)
        if tok.strip() and _THAI_RE.search(tok):
            out.append(_ZWSP)
    result = "".join(out)
    # Drop a trailing/space-adjacent ZWSP so it never sits before a real space.
    return re.sub(r"​(?=\s|$)", "", result)


def _apply_thai_font(run, size: Optional[int] = None, bold: bool = False) -> None:
    """Force a Thai-capable font + th-TH language on a run (and complex-script size).

    Thai is complex-script text, so bold must also be set via ``w:bCs`` – plain
    ``w:b`` alone is ignored by Word/LibreOffice for Thai glyphs.
    """
    run.font.name = THAI_FONT
    rpr = run._element.get_or_add_rPr()
    if bold and rpr.find(qn("w:bCs")) is None:
        rpr.append(rpr.makeelement(qn("w:bCs"), {}))
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), THAI_FONT)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = rpr.makeelement(qn("w:lang"), {})
        rpr.append(lang)
    for attr in ("w:val", "w:eastAsia", "w:bidi"):
        lang.set(qn(attr), "th-TH")
    if size is not None:
        szcs = rpr.find(qn("w:szCs"))
        if szcs is None:
            szcs = rpr.makeelement(qn("w:szCs"), {})
            rpr.append(szcs)
        szcs.set(qn("w:val"), str(int(size) * 2))  # half-points


def _configure_thai_doc(doc: Document) -> None:
    """Apply document-level settings that enable proper Thai line breaking."""
    settings = doc.settings.element
    csc = settings.find(qn("w:characterSpacingControl"))
    if csc is None:
        csc = settings.makeelement(qn("w:characterSpacingControl"), {})
        settings.insert(0, csc)
    csc.set(qn("w:val"), "doNotCompress")
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = settings.makeelement(qn("w:compat"), {})
        settings.append(compat)
    if compat.find(qn("w:applyBreakingRules")) is None:
        compat.insert(0, compat.makeelement(qn("w:applyBreakingRules"), {}))
    tfl = settings.find(qn("w:themeFontLang"))
    if tfl is None:
        tfl = settings.makeelement(qn("w:themeFontLang"), {})
        settings.append(tfl)
    for attr in ("w:val", "w:eastAsia", "w:bidi"):
        tfl.set(qn(attr), "th-TH")


def _set_base_style(doc: Document) -> None:
    # A4 paper. Margins: top/bottom 1", left/right 0.5", no gutter. Every fixed
    # table is scaled (in ``_set_fixed_col_widths``) to never exceed the content
    # width these margins leave.
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.gutter = Inches(0)
    style = doc.styles["Normal"]
    style.font.name = THAI_FONT
    style.font.size = Pt(14)
    # Content text: no extra spacing above/below paragraphs.
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), THAI_FONT)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = rpr.makeelement(qn("w:lang"), {})
        rpr.append(lang)
    for attr in ("w:val", "w:eastAsia", "w:bidi"):
        lang.set(qn(attr), "th-TH")
    _configure_thai_doc(doc)


def _add_run(paragraph, text: str, *, bold: bool = False, size: int = 14,
             color: Optional[RGBColor] = None):
    """Add a run, rendering embedded ``\\n`` as real Word line breaks.

    Thai text is segmented with U+200B so Word wraps it at word boundaries.
    """
    text = "" if text is None else _xml_safe(str(text))
    # Normalise line endings first (textarea input is often \r\n), then collapse
    # blank lines so multi-item content (PLOs/CLOs, etc.) renders as consecutive
    # lines with no empty gap between entries. Single line breaks – used
    # intentionally for layout (e.g. "\nวิชาเอก") – are preserved.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n[ \t]*(?:\n[ \t]*)+", "\n", text)
    parts = text.split("\n")
    run = None
    for idx, part in enumerate(parts):
        if idx > 0:
            if run is None:
                run = paragraph.add_run("")
                _apply_thai_font(run, size, bold)
            run.add_break()
        run = paragraph.add_run(_segment_thai(part))
        run.bold = bold
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        _apply_thai_font(run, size, bold)
    if run is None:
        run = paragraph.add_run("")
        run.bold = bold
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        _apply_thai_font(run, size, bold)
    return run


def _add_title(doc: Document, text: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, text, bold=True, size=20, color=_HEADING_COLOR)
    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(sp, subtitle, size=14, color=_MUTED_COLOR)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    # Black heading: 8pt above, none below.
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    # Keep the heading on the same page as the table/content that follows it.
    p.paragraph_format.keep_with_next = True
    _add_run(p, text, bold=True, size=16, color=_BLACK)


def _add_black_heading(doc: Document, text: str):
    """Inline bold black sub-heading: 8pt above, none below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    _add_run(p, text, bold=True, color=_BLACK)
    return p


def _add_field(doc: Document, label: str, value: Any) -> None:
    p = doc.add_paragraph()
    _add_run(p, f"{label}: ", bold=True)
    _add_run(p, _DASH if value in (None, "") else str(value))


def _add_paragraph_field(doc: Document, label: str, value: Any,
                         *, indent: float = 0) -> None:
    lp = doc.add_paragraph()
    _add_run(lp, label, bold=True)
    bp = doc.add_paragraph()
    if indent:
        # Indent the body so sub-items sit beneath (เยื้อง) the numbered heading.
        bp.paragraph_format.left_indent = Inches(indent)
    text = _DASH if value in (None, "") else str(value)
    _add_run(bp, text)


def _set_table_thai(table, margin_dxa: int = 108) -> None:
    """Set cell margins (≥108 DXA per spec) so wrapped Thai text isn't clipped."""
    tbl_pr = table._tbl.tblPr
    if tbl_pr.find(qn("w:tblCellMar")) is not None:
        return
    cell_mar = tbl_pr.makeelement(qn("w:tblCellMar"), {})
    for side in ("top", "left", "bottom", "right"):
        el = cell_mar.makeelement(qn(f"w:{side}"), {})
        el.set(qn("w:w"), str(margin_dxa))
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
    tbl_pr.append(cell_mar)


def _repeat_header_rows(table, header_rows: int = 1) -> None:
    """Repeat the top header row(s) on each page and stop rows tearing mid-page.

    Use for genuinely long tables that are expected to span pages.
    """
    for idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(tr_pr.makeelement(qn("w:cantSplit"), {}))
        if idx < header_rows and tr_pr.find(qn("w:tblHeader")) is None:
            el = tr_pr.makeelement(qn("w:tblHeader"), {})
            el.set(qn("w:val"), "true")
            tr_pr.append(el)


def _bind_header_to_body(table, header_rows: int = 1) -> None:
    """Keep the header row(s) with the first body row so a header is never left
    stranded at the foot of a page. Combined with the section heading's own
    keep-with-next, this means heading + header + first content row move to the
    next page together when they don't fit – i.e. the table's heading starts on a
    new page rather than printing a header with no content beneath it."""
    if len(table.rows) <= header_rows:
        return
    for row in table.rows[:header_rows]:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.keep_with_next = True


def _set_table_indent(table, inches: float) -> None:
    """Indent a whole table from the left margin by ``inches`` (via ``w:tblInd``)."""
    tbl_pr = table._tbl.tblPr
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = tbl_pr.makeelement(qn("w:tblInd"), {})
        tbl_pr.append(ind)
    ind.set(qn("w:w"), str(int(inches * 1440)))  # 1 inch = 1440 twips
    ind.set(qn("w:type"), "dxa")


def _keep_table_together(table) -> None:
    """Keep a whole (short) table on one page. If it doesn't fit, the heading +
    table move to the next page together (the heading uses keep-with-next)."""
    rows = table.rows
    last = len(rows) - 1
    for idx, row in enumerate(rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(tr_pr.makeelement(qn("w:cantSplit"), {}))
        if idx < last:  # bind each row to the next so the table can't split
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.keep_with_next = True


def _set_fixed_col_widths(table, widths_in: List[float]) -> None:
    """Pin column widths (inches) so Word honours them instead of auto-balancing.

    Sets a fixed table layout, a concrete total ``tblW`` (Word ignores per-column
    widths when ``tblW`` is ``auto``), the ``tblGrid`` columns, and every cell's
    width – merged cells get the sum of the grid columns they span.
    """
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = tbl_pr.makeelement(qn("w:tblLayout"), {})
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    # Scale down proportionally if the total exceeds the printable content width
    # (A4 minus 0.5" left/right margins) so no table ever runs past the margin.
    total_in = sum(widths_in)
    if total_in > _CONTENT_WIDTH_IN:
        scale = _CONTENT_WIDTH_IN / total_in
        widths_in = [w * scale for w in widths_in]

    emus = [int(Inches(w)) for w in widths_in]
    total_twips = sum(emus) // 635  # 1 twip = 635 EMU
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = tbl_pr.makeelement(qn("w:tblW"), {})
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))

    for col, emu in zip(table.columns, emus):
        col.width = Emu(emu)
    # Per-cell widths via the raw <w:tc> elements so gridSpan (merged) cells get
    # the correct spanned width and python-docx keeps tcW in the right XML order.
    for tr in table._tbl.tr_lst:
        idx = 0
        for tc in tr.tc_lst:
            span = tc.grid_span
            _Cell(tc, table).width = Emu(sum(emus[idx:idx + span]))
            idx += span


def _add_table(doc: Document, headers: List[str], rows: List[List[Any]],
               *, min_rows: int = 0, total_row: Optional[List[Any]] = None,
               blank_empty: bool = False,
               col_widths: Optional[List[float]] = None,
               center_cols: Optional[List[int]] = None,
               long_table: bool = False) -> None:
    """Render a bordered table.

    ``blank_empty`` leaves empty cells blank (strict-form look) instead of a dash.
    ``min_rows`` pads the body with empty rows so the form structure is preserved
    when there is little/no data. ``total_row`` appends one final summary row.
    ``col_widths`` (inches) pins fixed column widths instead of auto-fitting.
    ``center_cols`` (0-based column indices) centre-aligns those body cells.
    Content flows continuously down the page: the header row is always repeated on
    each page and kept with the first body row, so a header never prints alone at
    the foot of a page (the heading + header + first row move to the next page
    together). ``long_table`` is accepted for backward compatibility and no longer
    changes behaviour.
    """
    if not rows and not min_rows and not blank_empty:
        p = doc.add_paragraph()
        _add_run(p, "(ยังไม่มีข้อมูล)", color=_MUTED_COLOR)
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_thai(table)
    for cell, header in zip(table.rows[0].cells, headers):
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(cell.paragraphs[0], header, bold=True, size=13)
    body = list(rows)
    while len(body) < min_rows:
        body.append([""] * len(headers))
    if total_row is not None:
        body.append(total_row)
    empty_text = "" if blank_empty else _DASH
    centered = set(center_cols or [])
    for row in body:
        cells = table.add_row().cells
        for idx, (cell, value) in enumerate(zip(cells, row)):
            cell.paragraphs[0].clear()
            if idx in centered:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = empty_text if value in (None, "") else str(value)
            _add_run(cell.paragraphs[0], text, size=13)
    if col_widths:
        _set_fixed_col_widths(table, col_widths)
    # Flow continuously down the page: repeat the header on every page and keep it
    # with the first body row so the heading never strands a lone header.
    _repeat_header_rows(table, header_rows=1)
    _bind_header_to_body(table)


def _clear_cell_borders(cell) -> None:
    """Remove all borders on a cell (used for the spacer column between blocks)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    for side in ("top", "left", "bottom", "right"):
        el = borders.makeelement(qn(f"w:{side}"), {})
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tc_pr.append(borders)


def _add_teaching_method_table(doc: Document, general: Dict[str, Any]) -> None:
    """วิธีจัดการเรียนการสอน – two side-by-side blocks (วิธีการสอน | ร้อยละ) per format."""
    other_detail = _g(general, "other_detail", default="")
    other_label = "อื่นๆ (ระบุ) " + (
        str(other_detail) if other_detail not in (None, "", _DASH) else "..................")
    left = [("การบรรยาย", _g(general, "pct_lecture", default="")),
            ("การบรรยายเชิงอภิปราย", _g(general, "pct_discussion", default="")),
            ("กรณีศึกษา", _g(general, "pct_case", default=""))]
    right = [("การฝึกปฏิบัติ", _g(general, "pct_practice", default="")),
             ("กิจกรรมกลุ่ม", _g(general, "pct_group", default="")),
             (other_label, _g(general, "pct_other", default=""))]

    def _used(val: Any) -> bool:
        return str(val).strip() not in ("", _DASH)

    def _methods(items):
        return "\n".join(("■ " if _used(v) else "□ ") + lbl for lbl, v in items)

    def _pcts(items):
        return "\n".join((str(v) if _used(v) else "") for _lbl, v in items)

    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_thai(table)
    for col, txt in [(0, "วิธีการสอน"), (1, "ร้อยละของเวลาทั้งหมด"),
                     (3, "วิธีการสอน"), (4, "ร้อยละของเวลาทั้งหมด")]:
        cell = table.cell(0, col)
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(cell.paragraphs[0], txt, bold=True, size=13)
    data = table.rows[1].cells
    _add_run(data[0].paragraphs[0], _methods(left), size=13)
    data[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(data[1].paragraphs[0], _pcts(left), size=13)
    _add_run(data[3].paragraphs[0], _methods(right), size=13)
    data[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(data[4].paragraphs[0], _pcts(right), size=13)
    for row in range(2):  # borderless spacer column between the two blocks
        _clear_cell_borders(table.cell(row, 2))
    _set_fixed_col_widths(table, [2.45, 1.05, 0.25, 2.45, 1.05])
    _repeat_header_rows(table, header_rows=1)
    _bind_header_to_body(table, header_rows=1)


def _add_rubric_table(doc: Document, rows: List[List[Any]], *, min_rows: int = 2,
                      col_widths: Optional[List[float]] = None) -> None:
    """Rubric table with the official merged header.

    Layout (per format_tqf_new): ``ประเด็นการประเมิน`` spans the two header rows,
    and ``ระดับการประเมิน`` spans the five level columns above the ``5 4 3 2 1`` row.
    """
    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_thai(table)

    label = table.cell(0, 0).merge(table.cell(1, 0))
    label.paragraphs[0].clear()
    label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(label.paragraphs[0], "ประเด็นการประเมิน", bold=True, size=13)

    level_head = table.cell(0, 1).merge(table.cell(0, 5))
    level_head.paragraphs[0].clear()
    level_head.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(level_head.paragraphs[0], "ระดับการประเมิน", bold=True, size=13)

    for col, lvl in enumerate(["5", "4", "3", "2", "1"], start=1):
        cell = table.cell(1, col)
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(cell.paragraphs[0], lvl, bold=True, size=13)

    body = list(rows)
    while len(body) < min_rows:
        body.append([""] * 6)
    for row in body:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.paragraphs[0].clear()
            text = "" if value in (None, "") else str(value)
            _add_run(cell.paragraphs[0], text, size=13)
    if col_widths:
        _set_fixed_col_widths(table, col_widths)
    # Repeat the 2-row header on each page and keep it with the first body row so
    # the header is never stranded above a page break (heading + header + first row
    # move to the next page together when they don't fit).
    _repeat_header_rows(table, header_rows=2)
    _bind_header_to_body(table, header_rows=2)


def _checkbox(options: List[tuple], selected: Any) -> str:
    """Render ``■/□`` choice text (plain geometric squares, not emoji boxes).

    ``options`` items are (label, *match-keywords). Single-select: at most one box
    is ticked. An exact match wins first so a value like ``"วิชาเลือกเสรี"`` does
    not also tick ``"วิชาเลือก"`` (whose Thai text is a substring of it).
    """
    sel = str(selected or "").strip().lower()
    # Match keys per option: the label's Thai prefix (before " ("), the full
    # label, and any aliases.
    keysets = []
    for label, *aliases in options:
        prefix = label.split(" (")[0].strip()
        keysets.append([k.lower() for k in (prefix, label, *aliases) if k])

    checked = -1
    if sel:
        for i, keys in enumerate(keysets):
            if any(k == sel for k in keys):
                checked = i
                break
        else:  # legacy fuzzy: longest key that is a substring of the value
            best = 0
            for i, keys in enumerate(keysets):
                for k in keys:
                    if k in sel and len(k) > best:
                        checked, best = i, len(k)

    return "    ".join(("■ " if i == checked else "□ ") + options[i][0]
                       for i in range(len(options)))


def _fill_cell(cell, segments: List[tuple], *, size: int = 13) -> None:
    """Write ``(text, bold)`` segments into one cell paragraph."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    for text, bold in segments:
        _add_run(p, text, bold=bold, size=size)


def _merge_full_row(table, row_idx: int):
    """Merge every cell of a row into one full-width cell and return it."""
    cells = table.rows[row_idx].cells
    merged = cells[0]
    for cell in cells[1:]:
        merged = merged.merge(cell)
    return merged


def _add_footer_note(doc: Document, text: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    _add_run(p, text, size=12, color=_MUTED_COLOR)


def _g(data: Dict[str, Any], *keys: str, default: Any = _DASH) -> Any:
    """First non-empty value among ``keys`` (ctx fallbacks already merged in)."""
    for key in keys:
        value = data.get(key)
        if value not in (None, ""):
            return value
    return default


def _arr(data: Dict[str, Any], *keys: str) -> List[Any]:
    for key in keys:
        value = data.get(key)
        if value:
            return value
    return []


def _at(seq: List[Any], idx: int, default: Any = _DASH) -> Any:
    if idx < len(seq):
        value = seq[idx]
        return _DASH if value in (None, "") else value
    return default


def _fmt_date(value: Any) -> str:
    """Render a date as ``วัน/เดือน/ปี`` (dd/mm/yyyy).

    Accepts ISO ``yyyy-mm-dd`` (optionally with a time part); anything that does
    not parse is returned unchanged.
    """
    s = "" if value in (None, "", _DASH) else str(value).strip()
    m = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})", s)
    if not m:
        return s
    year, month, day = m.group(1), m.group(2), m.group(3)
    return f"{int(day):02d}/{int(month):02d}/{year}"


def _clo_no(text: Any, num: int) -> str:
    """Prefix a single CLO entry with its item number (``1.``, ``2.`` …).

    Empty/placeholder values are returned unchanged so padded rows stay blank.
    A value that already starts with that number (e.g. ``1.``/``1)``) is left
    as-is to avoid double numbering.
    """
    s = "" if text in (None, "", _DASH) else str(text).strip()
    if not s:
        return s
    if re.match(rf"^{num}\s*[.)]", s):
        return s
    return f"{num}. {s}"


def _clo_lines(text: Any) -> str:
    """Number each non-empty line of a CLO block as ``1.``, ``2.`` … (used for the
    free-text CLO cells that list several CLOs in one field)."""
    s = "" if text in (None, "", _DASH) else str(text)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    out, n = [], 0
    for line in s.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        n += 1
        if re.match(rf"^{n}\s*[.)]", stripped):
            out.append(stripped)
        else:
            out.append(f"{n}. {stripped}")
    return "\n".join(out)


def _max_indexed(data: Dict[str, Any], prefix: str,
                 exclude_prefix: Optional[str] = None) -> int:
    biggest = 0
    for key in data.keys():
        if not key.startswith(prefix):
            continue
        if exclude_prefix and key.startswith(exclude_prefix):
            continue
        suffix = key.rsplit("_", 1)[-1]
        if suffix.isdigit():
            biggest = max(biggest, int(suffix))
    return biggest


def _signature_image(value: Any) -> Optional[bytes]:
    """Decode a ``data:image/...;base64,...`` signature URL to raw bytes."""
    if not value or not isinstance(value, str):
        return None
    m = re.match(r"^data:image/[a-zA-Z0-9.+-]+;base64,(.+)$", value.strip(), re.DOTALL)
    if not m:
        return None
    try:
        return base64.b64decode(m.group(1))
    except (ValueError, base64.binascii.Error):
        return None


def _add_signatures(doc: Document, roles: List[Any]) -> None:
    """Append signature lines (name + date) for the given roles.

    Each item may be a plain role string, a ``(role, name)`` tuple, or a
    ``(role, name, signature)`` tuple. ``signature`` is a base64 image data URL;
    when present it is drawn over the line in place of the dotted placeholder.
    The name is printed inside the parentheses below the line.

    Laid out as a borderless 2-row, 3-column table per role: the top row holds
    (label | signature line/image | date) vertically centred together, and the
    parenthesised name sits on the row below, centred under the signature.
    """
    doc.add_paragraph()
    for role in roles:
        if isinstance(role, (list, tuple)):
            role_label = role[0]
            name = role[1] if len(role) > 1 else ""
            signature = role[2] if len(role) > 2 else ""
        else:
            role_label, name, signature = role, "", ""
        inner = str(name).strip() if name not in (None, "", _DASH) else ""

        table = doc.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Indent the whole block 0.5" from the left; columns shrunk so the block
        # (indent + columns) still stays within the right margin.
        _set_fixed_col_widths(table, [2.5, 2.5, 1.7])
        _set_table_indent(table, 0.5)
        # Borderless: clear every cell's borders.
        for row in table.rows:
            for cell in row.cells:
                _clear_cell_borders(cell)
        # Top row: label / signature / date, vertically centred as one line.
        label_cell, sign_cell, date_cell = table.rows[0].cells
        for cell in (label_cell, sign_cell, date_cell):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        lp = label_cell.paragraphs[0]
        lp.clear()
        _add_run(lp, f"ลงชื่อ {role_label}")

        # Signature: saved image if available, otherwise the dotted line.
        sp = sign_cell.paragraphs[0]
        sp.clear()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img = _signature_image(signature)
        if img is not None:
            try:
                sp.add_run().add_picture(io.BytesIO(img), width=Inches(1.6))
            except Exception:  # corrupt/unsupported image → fall back to the line
                _add_run(sp, "..............................................")
        else:
            _add_run(sp, "..............................................")

        dp = date_cell.paragraphs[0]
        dp.clear()
        today = date.today()
        _add_run(dp, f"วันที่ {today.day:02d}/{today.month:02d}/{today.year + 543}")

        # Bottom row: the parenthesised name centred under the signature.
        nm = table.cell(1, 1).paragraphs[0]
        nm.clear()
        nm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(nm, f"( {inner} )" if inner else "(                              )")

        # Keep each signature block intact so it never tears across a page break.
        _keep_table_together(table)
        doc.add_paragraph()


def _finalize(doc: Document) -> io.BytesIO:
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_tqf3_docx(general: Dict[str, Any], ctx: Dict[str, Any]) -> io.BytesIO:
    """Render a TQF3 (มคอ.3) Word document. ``general`` is ``tqf3.general_info``.

    ``ctx`` holds resolved fallbacks (faculty/program/course/term/instructor).
    """
    general = dict(general or {})
    doc = Document()
    _set_base_style(doc)

    course_code = _g(general, "course_code", default=ctx.get("course_code", _DASH))
    course_name = _g(general, "course_name", default=ctx.get("course_name", _DASH))
    _add_title(doc, "ประมวลการสอนและแผนการจัดการเรียนรู้ (มคอ.3)",
               "มหาวิทยาลัยราชภัฏเทพสตรี")
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(cp, f"{course_code} {course_name}".strip(), bold=True)

    # ---- ตารางข้อมูลรายวิชา (Table 0 ในแบบฟอร์มทางการ) ----
    credits = _g(general, "credits", default=ctx.get("credits", _DASH))
    info = doc.add_table(rows=14, cols=3)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_thai(info)
    # แถว 0: รหัสวิชา | ชื่อวิชา | จำนวนหน่วยกิต (สามคอลัมน์)
    head = info.rows[0].cells
    _fill_cell(head[0], [("รหัสวิชา ", True), (course_code, False)])
    _fill_cell(head[1], [("ชื่อวิชา ", True), (course_name, False)])
    _fill_cell(head[2], [("จำนวน ", True), (str(credits), False), (" หน่วยกิต", True)])
    status_line = _checkbox(
        [("วิชาบังคับ (Required)", "บังคับ", "required"),
         ("วิชาเลือก (Elective)", "เลือก", "elective"),
         ("วิชาเลือกเสรี (Free Elective)", "เสรี", "free")],
        _g(general, "course_status", default=""))
    stu_line = _checkbox(
        [("ภาคปกติ", "ปกติ"), ("ภาคพิเศษ", "พิเศษ"), ("บัณฑิตศึกษา", "บัณฑิต")],
        _g(general, "student_type", default=""))
    # ชั่วโมงให้คำปรึกษา + ช่องทางติดต่อ → "1 ชั่วโมง/สัปดาห์    ติดต่อ 062xxxxxxx"
    consult_hours = str(_g(general, "consult_hours", default="")).strip()
    consult_contact = str(_g(general, "hours_note", default="")).strip()
    consult_parts = []
    if consult_hours not in ("", _DASH):
        consult_parts.append((f"{consult_hours} ชั่วโมง/สัปดาห์", False))
    if consult_contact not in ("", _DASH):
        consult_parts.append((("    ติดต่อ " if consult_parts else "ติดต่อ ")
                              + consult_contact, False))
    if not consult_parts:
        consult_parts = [("", False)]
    rows_full = [
        [("สถานภาพของวิชา    ", True), (status_line, False)],
        ("split2",
         [("รายวิชาสังกัดคณะ\n", True),
          (_g(general, "faculty", default=ctx.get("faculty", _DASH)), False)],
         [("หลักสูตร ", True), (_g(general, "program", default=ctx.get("program", _DASH)), False),
          ("\nวิชาเอก ", True), (_g(general, "major", default=""), False)]),
        [("คำอธิบายรายวิชา ", True),
         (_g(general, "description", default=ctx.get("description", _DASH)), False)],
        [("รายวิชาที่บังคับเรียนก่อน (Pre-requisite) (ถ้ามี) ", True),
         (_g(general, "prereq", default=""), False)],
        [("ภาคการศึกษา ", True), (_g(general, "semester", default=ctx.get("semester", _DASH)), False),
         ("    ปีการศึกษา ", True),
         (_g(general, "academic_year", default=ctx.get("year", _DASH)), False),
         ("    ประเภทนักศึกษา ", True), (stu_line, False),
         ("    ชั้นปีที่ ", True), (_g(general, "year_level", default=""), False)],
        [("อาจารย์ผู้สอน ", True),
         (_g(general, "instructor", default=ctx.get("instructor", _DASH)), False),
         ("    ห้องพัก ", True), (_g(general, "office", default=""), False),
         ("    โทรศัพท์ ", True), (_g(general, "phone", default=""), False)],
        [("ห้องเรียน ", True), (_g(general, "location", "location_type", default=""), False),
         ("    ระบบออนไลน์ ", True), (_g(general, "online_system", default=""), False),
         ("\nกรณีเรียนภายนอกมหาวิทยาลัย (สถานประกอบการ/ชุมชน/อื่นๆ) ", True),
         (_g(general, "external_location", default=""), False)],
        [("วันที่จัดทำหรือปรับปรุงประมวลการสอนหรือแผนการจัดการเรียนรู้ครั้งล่าสุด ", True),
         (_fmt_date(_g(general, "last_updated", default="")), False)],
        [("ผลลัพธ์การเรียนรู้ ระดับหลักสูตร (Program Learning Outcomes: PLOs) "
          "(ให้ระบุเฉพาะ PLOs ที่เกี่ยวข้องกับรายวิชา และสอดคล้องกับเล่ม มคอ 2 หลักสูตร)\n", True),
         (str(_g(general, "plos", default="")).replace("—", "–"), False)],
        [("จุดประสงค์การเรียนรู้ระดับรายวิชา (Course Learning Outcomes: CLOs)\n", True),
         (_clo_lines(_g(general, "course_objective", "objectives", default="")), False)],
        [("จำนวนชั่วโมงที่ใช้ต่อสัปดาห์ในการจัดการเรียนรู้ (Hours/Week)\n", True),
         ("จำนวนชั่วโมงบรรยาย ", False), (_g(general, "lecture_hours", default="......"), False),
         (" ชั่วโมง    จำนวนชั่วโมงฝึกปฏิบัติ/ภาคสนาม/การฝึกงาน ", False),
         (_g(general, "lab_hours", default="......"), False),
         (" ชั่วโมง    จำนวนชั่วโมงการศึกษาด้วยตนเอง ", False),
         (_g(general, "self_hours", default="......"), False), (" ชั่วโมง", False)],
        [("จำนวนชั่วโมงต่อสัปดาห์ที่จะให้คำปรึกษาและแนะนำทางวิชาการแก่นักศึกษาเป็นรายบุคคล/"
          "ช่องทางการติดต่ออาจารย์ผู้สอน\n", True),
         *consult_parts],
        [("วิธีการ/ช่องทางสำหรับอุทธรณ์การเรียน กรณีการอุทธรณ์ผลคะแนนระหว่างภาคเรียนและคะแนนสอบ"
          "กลางภาค (ตามประกาศมหาวิทยาลัยราชภัฏเทพสตรี เรื่องระบบและแนวทางการอุทธรณ์ผลการเรียน"
          "การวัดและการประเมินผลการเรียนของนักศึกษา)\n", True),
         (_g(general, "appeal_channel", default=""), False)],
    ]
    for offset, row in enumerate(rows_full, start=1):
        if isinstance(row, tuple) and row and row[0] == "split2":
            cells = info.rows[offset].cells
            right = cells[1].merge(cells[2])
            _fill_cell(cells[0], row[1])
            _fill_cell(right, row[2])
        else:
            _fill_cell(_merge_full_row(info, offset), row)
    _repeat_header_rows(info, header_rows=1)
    _bind_header_to_body(info)

    # การพัฒนานักศึกษาตามผลลัพธ์การเรียนรู้ที่คาดหวัง (CLO table)
    _add_section_heading(doc, "การพัฒนานักศึกษาตามผลลัพธ์การเรียนรู้ที่คาดหวัง")
    clo_texts = _arr(general, "clo_text[]", "clo_desc[]")
    teachs = _arr(general, "teach_strategy[]")
    criterias = _arr(general, "assess_criteria[]")
    assesses = _arr(general, "assess_strategy[]")
    rows = [
        [_clo_no(_at(clo_texts, i), i + 1), _at(teachs, i), _at(criterias, i),
         _at(assesses, i)]
        for i in range(len(clo_texts))
    ]
    _add_table(doc, ["ผลลัพธ์การเรียนรู้ที่คาดหวังของรายวิชา (CLOs)",
                     "กลยุทธ์การสอนตาม CLOs", "เกณฑ์การวัดและการประเมินผล",
                     "วิธีการวัดและประเมินผลตาม CLOs"], rows, blank_empty=True)
    n1 = doc.add_paragraph()
    _add_run(n1, "หมายเหตุ : ให้ระบุรายละเอียดของ CLOs, PLOs และกลยุทธ์การสอน "
             "วิธีการวัดประเมินผลที่สอดคล้องตามเล่มหลักสูตรกำหนด ในกรณีหลักสูตรใช้เกณฑ์มาตรฐาน 2558 "
             "ให้นำ TQF เดิมที่กำหนดในหลักสูตร (มคอ 2) มาใช้ไปพรางก่อน", size=13, color=_MUTED_COLOR)

    # แผนการจัดการเรียนรู้ (weekly plan; activity + media share one column per format)
    _add_section_heading(doc, "แผนการจัดการเรียนรู้")
    weeks = _arr(general, "week[]")
    topics = _arr(general, "topic[]", "plan_topic[]")
    week_clos = _arr(general, "week_clo[]", "plan_clo[]")
    hours = _arr(general, "hours[]")
    activities = _arr(general, "activities[]")
    medias = _arr(general, "media[]", "plan_media[]")
    teachers = _arr(general, "teacher[]")
    plan_lec = _arr(general, "plan_lecture[]")
    plan_prac = _arr(general, "plan_practice[]")
    plan_rows = []
    hours_total = 0.0
    has_hours = False
    for i in range(len(topics)):
        wk = _at(weeks, i, default="")
        if str(wk).strip() == "":
            wk = i + 1
        hr = _at(hours, i, default="")
        if str(hr).strip() == "" and (plan_lec or plan_prac):
            try:
                total = float(_at(plan_lec, i, 0) or 0) + float(_at(plan_prac, i, 0) or 0)
            except (TypeError, ValueError):
                total = 0
            hr = total if total > 0 else ""
        try:
            hours_total += float(str(hr).strip())
            has_hours = True
        except (TypeError, ValueError):
            pass
        act = _at(activities, i, default="")
        med = _at(medias, i, default="")
        act = "" if act in (None, _DASH) else str(act).strip()
        med = "" if med in (None, _DASH) else str(med).strip()
        if act and med:
            act_media = f"{act}\nสื่อที่ใช้: {med}"
        elif med:
            act_media = f"สื่อที่ใช้: {med}"
        else:
            act_media = act
        plan_rows.append([
            wk, _at(topics, i, default=""), _at(week_clos, i, default=""),
            hr if str(hr).strip() else "",
            act_media, _at(teachers, i, default=""),
        ])
    total_hr = (str(int(hours_total)) if has_hours and hours_total.is_integer()
                else (str(hours_total) if has_hours else ""))
    _add_table(doc, ["สัปดาห์ที่", "หัวข้อ/รายละเอียด", "Lesson Learning Outcome : LLOs",
                     "จำนวนชั่วโมง", "กิจกรรมการเรียนการสอน สื่อที่ใช้ (ถ้ามี)", "ผู้สอน"],
               plan_rows, blank_empty=True,
               total_row=["", "", "รวม", total_hr, "", ""],
               col_widths=[0.65, 1.6, 1.5, 0.55, 2.2, 0.76],
               center_cols=[0, 3, 5], long_table=True)

    # วิธีจัดการเรียนการสอน (สองบล็อกเรียงข้างกัน ตามแบบฟอร์ม)
    _add_section_heading(doc, "วิธีจัดการเรียนการสอน")
    _add_teaching_method_table(doc, general)

    # แผนการประเมินตามผลลัพธ์การเรียนรู้ที่คาดหวังของรายวิชา
    _add_section_heading(doc, "แผนการประเมินตามผลลัพธ์การเรียนรู้ที่คาดหวังของรายวิชา")
    a_clo = _arr(general, "assess_clo[]")
    a_act = _arr(general, "assess_activity[]", "assess_method[]")
    a_crit = _arr(general, "assess_plan_criteria[]")
    a_pct = _arr(general, "assess_pct[]", "assess_ratio[]")
    n_assess = max(len(a_clo), len(a_act), len(a_crit), len(a_pct))
    assess_rows = [
        [_clo_no(_at(a_clo, i, default=""), i + 1), _at(a_act, i, default=""),
         _at(a_crit, i, default=""), _at(a_pct, i, default="")]
        for i in range(n_assess)
    ]
    _add_table(doc, ["ผลลัพธ์การเรียนรู้ที่คาดหวังของรายวิชา (CLOs)",
                     "กิจกรรมการจัดการเรียนรู้ของผู้เรียน",
                     "เกณฑ์การประเมินผลลัพธ์การเรียนรู้ระดับรายวิชา",
                     "สัดส่วนของการประเมินผล"], assess_rows, blank_empty=True,
               col_widths=[1.12, 2.63, 2.64, 0.93], center_cols=[0, 3])
    n2 = doc.add_paragraph()
    _add_run(n2, "หมายเหตุ กรณีหลักสูตรใช้เกณฑ์มาตรฐาน 2558 ให้ระบุ CLOs "
             "ตามที่ปรับในหมวดการพัฒนานักศึกษาตามผลลัพธ์การเรียนรู้ที่คาดหวัง",
             size=13, color=_MUTED_COLOR)

    # เครื่องมือสำคัญที่ใช้ประเมินผลลัพธ์การเรียนรู้ (Rubric Score)
    _add_section_heading(doc, "เครื่องมือสำคัญที่ใช้ประเมินผลลัพธ์การเรียนรู้ "
                         "(Rubric Score และเกณฑ์การตัดสินการบรรลุผลลัพธ์การเรียนรู้)")
    r_topic = _arr(general, "rubric_topic[]")
    r_l5 = _arr(general, "rubric_l5[]")
    r_l4 = _arr(general, "rubric_l4[]")
    r_l3 = _arr(general, "rubric_l3[]")
    r_l2 = _arr(general, "rubric_l2[]")
    r_l1 = _arr(general, "rubric_l1[]")
    rubric_rows = [
        [_at(r_topic, i), _at(r_l5, i), _at(r_l4, i), _at(r_l3, i), _at(r_l2, i), _at(r_l1, i)]
        for i in range(len(r_topic))
        if str(_at(r_topic, i, default="")).strip()
    ]
    _add_rubric_table(doc, rubric_rows, min_rows=0,
                      col_widths=[2.07, 1.13, 1.13, 1.13, 1.13, 1.13])
    n3 = doc.add_paragraph()
    _add_run(n3, "ควรแนบ Rubric Score ที่ใช้ประเมินงานที่สะท้อน CLOs โดยประเด็นการประเมิน"
             "ควรวัดพฤติกรรมบ่งชี้ ความรู้ ทักษะ จริยธรรม หรือคุณลักษณะที่ต้องการวัดให้ชัดเจน "
             "และควรมีความสอดคล้องและสะท้อนการบรรลุ CLOs", size=13, color=_MUTED_COLOR)
    n4 = doc.add_paragraph()
    _add_run(n4, "หลักสูตรต้องกำหนด Rubric Score สำหรับใช้ประเมิน 1. ทักษะสื่อสาร "
             "2. ทักษะการทำงานร่วมกัน 3. ทักษะด้านเทคโนโลยีดิจิทัล "
             "4. ทักษะการคิดอย่างมีวิจารณญาณและการแก้ปัญหา 5. ความคิดสร้างสรรค์ "
             "6. จิตสำนึกสาธารณะ ซึ่งผู้สอนสามารถนำมาปรับประยุกต์ใช้ได้อย่างเหมาะสม",
             size=13, color=_MUTED_COLOR)

    # ตำราและเอกสารที่ใช้ประกอบการเรียนการสอน (APA)
    _add_section_heading(doc, "ตำราและเอกสารที่ใช้ประกอบการเรียนการสอน "
                         "(เขียนตามแบบบรรณานุกรม APA)")
    refs = _g(general, "references", default="")
    rp = doc.add_paragraph()
    _add_run(rp, str(refs) if refs not in (None, "", _DASH) else
             ("." * 100 + "\n" + "." * 100 + "\n" + "." * 100))

    # การประเมินและปรับปรุงการจัดการเรียนรู้ของรายวิชา
    _add_section_heading(doc, "การประเมินและปรับปรุงการจัดการเรียนรู้ของรายวิชา")
    _add_paragraph_field(doc, "1.  กลยุทธ์การประเมินการจัดการเรียนการสอนของรายวิชา",
                         _g(general, "course_eval_strategy"), indent=0.4)
    _add_paragraph_field(doc, "2.  การปรับปรุงการจัดการเรียนรู้ของรายวิชา",
                         _g(general, "course_improve", "improvement_strategy"), indent=0.4)
    g2 = doc.add_paragraph()
    _add_run(g2, "(อธิบายกลไกและวิธีการปรับปรุงการสอน การวัดและการประเมินผลลัพธ์การเรียนรู้ "
             "เช่น คณะ/สาขาวิชา มีการกำหนดกลไกและวิธีการปรับปรุงการสอนไว้อย่างไร การวิจัยในชั้นเรียน "
             "การประชุมเชิงปฏิบัติการเพื่อพัฒนาการเรียนการสอน หรือการประชุมคณะกรรมการบริหารหลักสูตร "
             "เป็นต้น)", size=12, color=_MUTED_COLOR)
    _add_paragraph_field(doc, "3.  กระบวนการยืนยัน (verification) ผลสัมฤทธิ์ทางการเรียน "
                         "และผลลัพธ์การเรียนรู้ของนักศึกษา/ผู้เรียน",
                         _g(general, "verification"), indent=0.4)
    g3 = doc.add_paragraph()
    _add_run(g3, "(อธิบายกระบวนการที่ใช้ในการยืนยัน (verification) ผลสัมฤทธิ์ของนักศึกษาหรือ"
             "ผลลัพธ์การเรียนรู้ของรายวิชา โดยกระบวนการอาจจะต่างกันไปสำหรับรายวิชาที่แตกต่างกัน "
             "หรือเป็นไปตามมาตรฐานผลการเรียนรู้แต่ละด้านที่หลักสูตรกำหนด)", size=12, color=_MUTED_COLOR)

    sigs = ctx.get("signatures", {})
    inst_name = _g(general, "instructor", default=ctx.get("instructor", ""))
    # Chair name/signature are bound to the verified approver (ctx); the free-text
    # ``course_owner`` is only a printed-name fallback and never fetches an image.
    chair_name = ctx.get("chair_name") or _g(general, "course_owner", default="")
    _add_signatures(doc, [
        ("อาจารย์ผู้สอน", inst_name,
         sigs.get(str(inst_name).strip()) or ctx.get("instructor_signature", "")),
        ("ประธานบริหารหลักสูตร", chair_name, ctx.get("chair_signature", "")),
    ])
    _add_footer_note(doc, "หมายเหตุ Update มีนาคม 2569 ตามเกณฑ์มาตรฐานหลักสูตร 2565")

    return _finalize(doc)


def build_tqf4_docx(general: Dict[str, Any], ctx: Dict[str, Any]) -> io.BytesIO:
    """Render a TQF4 (มคอ.4) field-experience course spec. ``general`` is ``tqf4.general_info``."""
    general = dict(general or {})
    doc = Document()
    _set_base_style(doc)

    course_code = _g(general, "course_code", default=ctx.get("course_code", _DASH))
    course_name = _g(general, "course_name", default=ctx.get("course_name", _DASH))
    _add_title(doc, "ประมวลการสอนและแผนการจัดการเรียนรู้รายวิชาประสบการณ์ภาคสนาม (มคอ.4)",
               "มหาวิทยาลัยราชภัฏเทพสตรี")
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(cp, f"{course_code} {course_name}".strip(), bold=True)

    # 1) ข้อมูลรายวิชา
    _add_section_heading(doc, "1) ข้อมูลรายวิชา")
    _add_field(doc, "สถาบันการศึกษา", _g(general, "university",
               default="มหาวิทยาลัยราชภัฏเทพสตรี"))
    _add_field(doc, "คณะ", _g(general, "faculty", default=ctx.get("faculty", _DASH)))
    _add_field(doc, "หลักสูตร", _g(general, "program", default=ctx.get("program", _DASH)))
    _add_field(doc, "วิชาเอก", _g(general, "major"))
    _add_field(doc, "รหัสวิชา", course_code)
    _add_field(doc, "ชื่อวิชา", course_name)
    _add_field(doc, "สถานภาพของวิชา", _g(general, "course_status"))
    _add_field(doc, "หน่วยกิต", _g(general, "credits", default=ctx.get("credits", _DASH)))
    _add_field(doc, "ประเภทนักศึกษา", _g(general, "student_type"))
    _add_field(doc, "ชั้นปี", _g(general, "year_level"))
    _add_field(doc, "ภาคการศึกษา / ปีการศึกษา",
               f"{_g(general, 'semester', default=ctx.get('semester', _DASH))} / "
               f"{_g(general, 'academic_year', default=ctx.get('year', _DASH))}")
    _add_paragraph_field(doc, "คำอธิบายรายวิชา",
                         _g(general, "description", default=ctx.get("description", _DASH)))
    _add_field(doc, "รายวิชาที่เรียนก่อน", _g(general, "prereq"))
    _add_paragraph_field(doc, "เงื่อนไขที่สำคัญของการฝึกประสบการณ์", _g(general, "field_condition"))
    adv = _arr(general, "advisor[]")
    ofc = _arr(general, "office[]")
    tel = _arr(general, "phone[]")
    if not adv and general.get("advisor") not in (None, ""):
        adv, ofc, tel = [general.get("advisor")], [general.get("office")], [general.get("phone")]
    adv_rows = [[_at(adv, i), _at(ofc, i), _at(tel, i)] for i in range(len(adv))]
    _add_section_heading(doc, "อาจารย์ที่ปรึกษา/อาจารย์นิเทศการฝึกประสบการณ์ภาคสนาม")
    _add_table(doc, ["ชื่ออาจารย์ที่ปรึกษา/อาจารย์นิเทศ", "ห้องพัก", "โทรศัพท์"], adv_rows)
    _add_field(doc, "ห้องเรียน/สถานที่", _g(general, "location", "location_type"))
    _add_field(doc, "ระบบออนไลน์", _g(general, "online_system"))
    _add_field(doc, "กรณีฝึกภายนอกมหาวิทยาลัย", _g(general, "external_location"))

    # 2) จุดประสงค์ + ภาระงาน
    _add_section_heading(doc, "2) จุดประสงค์การจัดการเรียนรู้ของการฝึกประสบการณ์ภาคสนาม")
    _add_paragraph_field(doc, "จุดประสงค์การจัดการเรียนรู้",
                         _g(general, "field_objective", "objectives"))
    _add_field(doc, "ชั่วโมงบรรยาย", _g(general, "lecture_hours"))
    _add_field(doc, "ชั่วโมงเตรียมความพร้อมของนักศึกษา", _g(general, "prep_hours"))
    _add_field(doc, "ชั่วโมงฝึกปฏิบัติ/ภาคสนาม/ฝึกงาน", _g(general, "practice_hours", "lab_hours"))
    _add_field(doc, "ชั่วโมงการศึกษาด้วยตนเอง", _g(general, "self_hours"))
    _add_field(doc, "ชั่วโมงที่อาจารย์นิเทศให้คำแนะนำ", _g(general, "advisor_consult_hours", "consult_hours"))
    _add_paragraph_field(doc, "ช่องทางสำหรับอุทธรณ์การเรียน", _g(general, "appeal_channel"))
    _add_paragraph_field(doc, "การจัดการเรียนรู้รายวิชาการฝึกประสบการณ์ภาคสนาม",
                         _g(general, "field_management"))

    # 3) รายงาน/งานที่มอบหมาย
    _add_section_heading(doc, "3) รายงานหรืองานที่นักศึกษาได้รับมอบหมายและกำหนดการส่งงาน")
    rep = _arr(general, "report[]")
    rep_crit = _arr(general, "report_criteria[]")
    rep_rows = [[_at(rep, i), _at(rep_crit, i)] for i in range(len(rep))]
    _add_table(doc, ["รายงาน/งานที่มอบหมาย/กำหนดส่งงาน", "เกณฑ์การวัดประเมินผลลัพธ์การเรียนรู้"],
               rep_rows)
    _add_paragraph_field(doc, "การติดตามผลการเรียนรู้การฝึกประสบการณ์ภาคสนาม",
                         _g(general, "follow_up"))
    _add_paragraph_field(doc, "หน้าที่และความรับผิดชอบของพี่เลี้ยงในสถานประกอบการ",
                         _g(general, "mentor_duties"))
    _add_paragraph_field(doc, "หน้าที่และความรับผิดชอบของอาจารย์ที่ปรึกษา/อาจารย์นิเทศ",
                         _g(general, "advisor_duties"))
    _add_paragraph_field(doc, "สิ่งอำนวยความสะดวกและการสนับสนุนที่ต้องการจากสถานประกอบการ",
                         _g(general, "facilities"))

    # 4) การวางแผนและเตรียมการ
    _add_section_heading(doc, "4) การวางแผนและการเตรียมการสำหรับการฝึกประสบการณ์ภาคสนาม")
    _add_paragraph_field(doc, "การกำหนดสถานที่ฝึกประสบการณ์ภาคสนาม", _g(general, "prep_location"))
    _add_paragraph_field(doc, "การเตรียมอาจารย์ที่ปรึกษา/อาจารย์นิเทศ", _g(general, "prep_advisor"))
    _add_paragraph_field(doc, "การเตรียมความพร้อมนักศึกษา", _g(general, "prep_student"))
    _add_paragraph_field(doc, "การเตรียมพี่เลี้ยงในสถานประกอบการ", _g(general, "prep_mentor"))
    _add_paragraph_field(doc, "การจัดการความเสี่ยง", _g(general, "risk_mgmt"))

    # 5) การพัฒนานักศึกษาตามผลลัพธ์การเรียนรู้
    _add_section_heading(doc, "5) การพัฒนานักศึกษาตามผลลัพธ์การเรียนรู้ที่คาดหวังของหลักสูตร")
    _add_paragraph_field(doc, "ผลลัพธ์การเรียนรู้ระดับหลักสูตร (PLOs) ที่เกี่ยวข้อง",
                         str(_g(general, "plos")).replace("—", "–"))
    _add_paragraph_field(doc, "จุดประสงค์การเรียนรู้ระดับรายวิชา (Course Learning Outcomes: CLOs)",
                         _clo_lines(_g(general, "course_objective")))
    clo_texts = _arr(general, "clo_text[]", "clo_desc[]")
    clo_plo = _arr(general, "clo_plo[]", "plo[]")
    teachs = _arr(general, "teach_strategy[]")
    assesses = _arr(general, "assess_strategy[]")
    clo_rows = [
        [_clo_no(_at(clo_texts, i), i + 1), _at(clo_plo, i), _at(teachs, i),
         _at(assesses, i)]
        for i in range(len(clo_texts))
    ]
    _add_table(doc, ["CLOs ฝึกประสบการณ์ภาคสนาม", "PLOs ที่รับผิดชอบ",
                     "กลยุทธ์การฝึกตาม CLOs", "วิธีการวัดและประเมินผลตาม CLOs"], clo_rows)

    # 6) แผนการฝึกรายสัปดาห์
    _add_section_heading(doc, "6) แผนการฝึกประสบการณ์ภาคสนาม")
    weeks = _arr(general, "week[]")
    topics = _arr(general, "topic[]")
    week_clos = _arr(general, "week_clo[]")
    hours = _arr(general, "hours[]")
    activities = _arr(general, "activities[]")
    supervisors = _arr(general, "supervisor[]")
    notes = _arr(general, "note[]")
    plan_rows = []
    for i in range(len(topics)):
        wk = _at(weeks, i, default="")
        if str(wk).strip() == "":
            wk = i + 1
        plan_rows.append([
            wk, _at(topics, i), _at(week_clos, i), _at(hours, i),
            _at(activities, i), _at(supervisors, i), _at(notes, i),
        ])
    _add_table(doc, ["สัปดาห์", "หัวข้อ/รายละเอียด", "CLOs", "ชั่วโมง",
                     "กิจกรรมการฝึกภาคสนาม", "ผู้ดูแลกิจกรรม", "หมายเหตุ"], plan_rows,
               long_table=True)

    # 7) แผนการประเมิน
    _add_section_heading(doc, "7) แผนการประเมินตามผลลัพธ์การเรียนรู้ของรายวิชาฝึกประสบการณ์ภาคสนาม")
    a_clo = _arr(general, "assess_clo[]")
    a_act = _arr(general, "assess_activity[]")
    a_crit = _arr(general, "assess_plan_criteria[]")
    a_pct = _arr(general, "assess_pct[]")
    a_eval = _arr(general, "assess_evaluator[]")
    assess_rows = [
        [_clo_no(_at(a_clo, i), i + 1), _at(a_act, i), _at(a_crit, i), _at(a_pct, i),
         _at(a_eval, i)]
        for i in range(len(a_clo))
    ]
    _add_table(doc, ["CLOs", "กิจกรรมการประเมินผลการฝึก", "เกณฑ์การประเมิน",
                     "สัดส่วน (%)", "ผู้ประเมินผล"], assess_rows)

    # 8) Rubric Score
    _add_section_heading(doc, "8) เครื่องมือสำคัญที่ใช้ประเมินผลลัพธ์การเรียนรู้ (Rubric Score)")
    r_topic = _arr(general, "rubric_topic[]")
    r_l5 = _arr(general, "rubric_l5[]")
    r_l4 = _arr(general, "rubric_l4[]")
    r_l3 = _arr(general, "rubric_l3[]")
    r_l2 = _arr(general, "rubric_l2[]")
    r_l1 = _arr(general, "rubric_l1[]")
    rubric_rows = [
        [_at(r_topic, i), _at(r_l5, i), _at(r_l4, i), _at(r_l3, i), _at(r_l2, i), _at(r_l1, i)]
        for i in range(len(r_topic))
        if str(_at(r_topic, i, default="")).strip()
    ]
    _add_rubric_table(doc, rubric_rows, min_rows=0)

    # 9) การประเมินและปรับปรุงจากผู้เกี่ยวข้อง
    _add_section_heading(doc, "9) การประเมินและปรับปรุงการจัดการเรียนรู้จากผู้ที่เกี่ยวข้อง")
    _add_paragraph_field(doc, "นักศึกษาที่ฝึกประสบการณ์ภาคสนาม", _g(general, "eval_student"))
    _add_paragraph_field(doc, "พี่เลี้ยงในสถานประกอบการ", _g(general, "eval_mentor"))
    _add_paragraph_field(doc, "อาจารย์ที่ปรึกษา/อาจารย์นิเทศ", _g(general, "eval_advisor"))
    _add_paragraph_field(doc, "อื่นๆ เช่น บัณฑิตจบใหม่", _g(general, "eval_other"))
    _add_paragraph_field(doc, "กระบวนการยืนยัน (verification) ผลการประเมินและการวางแผนปรับปรุง",
                         _g(general, "verification"))

    sigs = ctx.get("signatures", {})
    inst_name = _g(general, "instructor", default=ctx.get("instructor", ""))
    # Chair name/signature are bound to the verified approver (ctx); the free-text
    # ``course_owner`` is only a printed-name fallback and never fetches an image.
    chair_name = ctx.get("chair_name") or _g(general, "course_owner", default="")
    _add_signatures(doc, [
        ("อาจารย์ผู้รับผิดชอบรายวิชา", inst_name,
         sigs.get(str(inst_name).strip()) or ctx.get("instructor_signature", "")),
        ("ประธานบริหารหลักสูตร", chair_name, ctx.get("chair_signature", "")),
    ])

    return _finalize(doc)


def build_tqf5_docx(data: Dict[str, Any], ctx: Dict[str, Any]) -> io.BytesIO:
    """Render a TQF5 (มคอ.5) Word document. ``data`` is ``tqf5.actual_teaching``."""
    data = dict(data or {})
    doc = Document()
    _set_base_style(doc)

    course_code = _g(data, "course_code", default=ctx.get("course_code", _DASH))
    course_name = _g(data, "course_name", default=ctx.get("course_name", _DASH))
    _add_title(doc, "รายงานผลการจัดการเรียนรู้ระดับรายวิชาของหลักสูตร",
               "มหาวิทยาลัยราชภัฏเทพสตรี")
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(cp, f"{course_code} {course_name}".strip(), bold=True)

    # ---- ตารางข้อมูลรายวิชา (Table 0 ในแบบฟอร์มทางการ มคอ.5) ----
    credits = _g(data, "credits", default=ctx.get("credits", _DASH))
    info = doc.add_table(rows=13, cols=3)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_thai(info)
    head = info.rows[0].cells
    _fill_cell(head[0], [("รหัสวิชา ", True), (course_code, False)])
    _fill_cell(head[1], [("ชื่อวิชา ", True), (course_name, False)])
    _fill_cell(head[2], [("จำนวน ", True), (str(credits), False), (" หน่วยกิต", True)])
    status_line = _checkbox(
        [("วิชาบังคับ (Required)", "บังคับ", "required"),
         ("วิชาเลือก (Elective)", "เลือก", "elective"),
         ("วิชาเลือกเสรี (Free Elective)", "เสรี", "free")],
        _g(data, "course_status", default=""))
    stu_line = _checkbox(
        [("ภาคปกติ", "ปกติ"), ("ภาคพิเศษ", "พิเศษ"), ("บัณฑิตศึกษา", "บัณฑิต")],
        _g(data, "student_type", default=""))
    rows_full = [
        [("สถานภาพของวิชา    ", True), (status_line, False)],
        ("split2",
         [("รายวิชาสังกัดคณะ\n", True),
          (_g(data, "faculty", default=ctx.get("faculty", _DASH)), False)],
         [("หลักสูตร ", True), (_g(data, "program", default=ctx.get("program", _DASH)), False),
          ("\nวิชาเอก ", True), (_g(data, "major", default=""), False)]),
        [("คำอธิบายรายวิชา ", True),
         (_g(data, "course_desc", default=ctx.get("description", _DASH)), False)],
        [("รายวิชาที่บังคับเรียนก่อน (Pre-requisite) (ถ้ามี) ", True),
         (_g(data, "prereq", default=""), False)],
        [("ผลลัพธ์การเรียนรู้ ระดับหลักสูตร (Program Learning Outcomes: PLOs) "
          "(ให้ระบุเฉพาะ PLOs ที่เกี่ยวข้องกับรายวิชา และสอดคล้องกับเล่ม มคอ 2 หลักสูตร)\n", True),
         (str(_g(data, "plos", default="")).replace("—", "–"), False)],
        [("จุดประสงค์การจัดการเรียนรู้ ระดับรายวิชา (Course Learning Outcomes: CLOs) "
          "(รวมทั้งรายวิชาฝึกประสบการณ์ภาคสนาม)\n", True),
         (_clo_lines(_g(data, "course_objective", default="")), False)],
        [("ภาคการศึกษา ", True), (_g(data, "semester", default=ctx.get("semester", _DASH)), False),
         ("    ปีการศึกษา ", True),
         (_g(data, "academic_year", default=ctx.get("year", _DASH)), False),
         ("    ประเภทนักศึกษา ", True), (stu_line, False),
         ("    ชั้นปีที่ ", True), (_g(data, "year_level", default=""), False)],
        [("อาจารย์ผู้สอน ", True),
         (_g(data, "instructors", default=ctx.get("instructor", _DASH)), False)],
        [("ห้องเรียน ", True), (_g(data, "location", "location_type", default=""), False)],
        [("รายงานจำนวนชั่วโมงที่สอนจริงและที่คลาดเคลื่อนในการจัดการเรียนรู้ (ถ้ามี) "
          "และแนวทางการจัดการแก้ไข\n", True),
         (_g(data, "teaching_hours", default=""), False)],
        [("รายงานหัวข้อที่สอนไม่ครอบคลุมจากแผนการจัดการเรียนรู้ที่กำหนดไว้ (ถ้ามี) "
          "และแนวทางการจัดการแก้ไข\n", True),
         (_g(data, "uncovered_topics", "deviations", default=""), False)],
        [("รายงานความสอดคล้องกับจุดประสงค์การเรียนรู้ และมีการปรับปรุงการจัดการเรียนรู้อย่างไร\n",
          True), (_g(data, "during_improve", default=""), False)],
    ]
    for offset, row in enumerate(rows_full, start=1):
        if isinstance(row, tuple) and row and row[0] == "split2":
            cells = info.rows[offset].cells
            right = cells[1].merge(cells[2])
            _fill_cell(cells[0], row[1])
            _fill_cell(right, row[2])
        else:
            _fill_cell(_merge_full_row(info, offset), row)
    _repeat_header_rows(info, header_rows=1)
    _bind_header_to_body(info)

    # การจัดการเรียนรู้และวิธีการประเมินผลที่ดำเนินการ (CLO table)
    _add_section_heading(doc, "การจัดการเรียนรู้และวิธีการประเมินผลที่ดำเนินการ"
                         "เพื่อทำให้เกิดผลลัพธ์การเรียนรู้ตามที่ระบุในรายละเอียดรายวิชา")
    legacy_codes = _arr(data, "clo_code[]")
    legacy_methods = _arr(data, "clo_method[]")
    legacy_assess = _arr(data, "clo_assess[]")
    legacy_results = _arr(data, "clo_result[]")
    clo_n = _max_indexed(data, "clo_desc_")
    if clo_n == 0 and legacy_codes:
        clo_n = len(legacy_codes)
    clo_rows = []
    for i in range(1, clo_n + 1):
        clo_rows.append([
            _clo_no(_g(data, f"clo_desc_{i}", default=_at(legacy_codes, i - 1)), i),
            _g(data, f"clo_teach_{i}", default=_at(legacy_methods, i - 1)),
            _g(data, f"clo_assess_{i}", default=_at(legacy_assess, i - 1)),
            _g(data, f"clo_result_{i}", default=_at(legacy_results, i - 1)),
            _g(data, f"clo_improve_{i}"),
        ])
    _add_table(doc, ["ผลลัพธ์การเรียนรู้ที่คาดหวังของรายวิชา (CLOs)",
                     "กลยุทธ์การสอน/วิธีการจัดการเรียนรู้ที่ได้ดำเนินการ",
                     "วิธีการประเมินผลที่ได้ดำเนินการ/เกณฑ์การวัดและการประเมินผล",
                     "ผลที่เกิดกับนักศึกษา (บรรลุผลลัพธ์การเรียนรู้ระดับรายวิชา/ระดับหลักสูตร)",
                     "แนวทางการพัฒนาปรับปรุง เพื่อให้นักศึกษาบรรลุตามแต่ละ CLOs และ PLOs ที่รับผิดชอบ"],
               clo_rows, long_table=True)
    note = doc.add_paragraph()
    _add_run(note, "หมายเหตุ กรณีรายวิชาฝึกประสบการณ์วิชาชีพ "
             "ให้คำนึงถึงผลลัพธ์การเรียนรู้ที่กำหนดใน มคอ.2", size=13, color=_MUTED_COLOR)

    # สรุปผลการจัดการเรียนการสอนของรายวิชา (numbered items 1–10)
    _add_section_heading(doc, "สรุปผลการจัดการเรียนการสอนของรายวิชา")
    _add_field(doc, "1. จำนวนนักศึกษาที่ลงทะเบียน",
               _g(data, "n_registered", "students_enrolled"))
    _add_field(doc, "2. จำนวนนักศึกษาที่คงอยู่เมื่อสิ้นสุดภาคการศึกษา",
               _g(data, "n_remain", "students_finished"))
    _add_field(doc, "3. จำนวนนักศึกษาที่ถอน (W)",
               _g(data, "n_withdraw", "students_withdrawn"))

    _add_black_heading(doc, "4. การกระจายของระดับคะแนน (เกรด) (แสดงแยกตามสาขา) (ถ้ามี)")
    grade_n = _max_indexed(data, "g_level_")
    grade_rows: List[List[Any]] = []
    if grade_n:
        for i in range(1, grade_n + 1):
            grade_rows.append([
                _g(data, f"g_level_{i}"),
                _g(data, f"g_count_{i}"),
                _g(data, f"g_percent_{i}"),
            ])
    else:
        for grade in ["A", "B+", "B", "C+", "C", "D+", "D", "F", "I", "S", "U", "W", "M"]:
            grade_rows.append([grade, data.get(f"grade_{grade}", 0), _DASH])
    _add_table(doc, ["ระดับคะแนน", "จำนวนนักศึกษา (คน)", "คิดเป็นร้อยละ"], grade_rows)

    _add_black_heading(doc, "5. การบรรลุผลลัพธ์การเรียนรู้ระดับรายวิชา "
                       "และระดับหลักสูตรตาม PLOs ที่รับผิดชอบ")
    _add_paragraph_field(doc, "5.1 การบรรลุผลลัพธ์การเรียนรู้ระดับรายวิชา (CLO) "
                         "เกณฑ์การวัดและการประเมินผลลัพธ์/ Rubric Score ที่กำหนด",
                         _g(data, "clo_achieve"))
    _add_paragraph_field(doc, "5.2 การบรรลุผลลัพธ์การเรียนรู้ระดับหลักสูตร (PLO ตัวที่รับผิดชอบ) "
                         "เกณฑ์การวัดและการประเมินผลลัพธ์/ Rubric Score ที่กำหนด",
                         _g(data, "plo_achieve"))
    _add_paragraph_field(doc, "6. ปัจจัยที่ทำให้นักศึกษาไม่บรรลุผลลัพธ์การเรียนรู้ "
                         "ระดับรายวิชา และระดับหลักสูตร",
                         _g(data, "fail_factors"))
    _add_paragraph_field(doc, "7. แนวทางการปรับปรุงแก้ไข กรณีที่นักศึกษาไม่บรรลุผลลัพธ์การเรียนรู้ "
                         "ระดับรายวิชา และระดับหลักสูตร",
                         _g(data, "fail_fix"))
    _add_paragraph_field(doc, "8. กระบวนการยืนยัน (verification) "
                         "ผลสัมฤทธิ์/ผลลัพธ์การเรียนรู้ของนักศึกษา",
                         _g(data, "verification", "verification_method"))

    _add_black_heading(doc, "9. ปัญหาและผลกระทบต่อการดำเนินการจัดการเรียนรู้")
    issue_n = _max_indexed(data, "issue_", exclude_prefix="issue_fix_")
    issue_rows = [
        [_g(data, f"issue_{i}"), _g(data, f"issue_fix_{i}")]
        for i in range(1, issue_n + 1)
    ]
    _add_table(doc, ["ประเด็น", "แนวทางการปรับแก้ไข"], issue_rows)

    _add_black_heading(doc, "10. การประเมินผลรายวิชา")
    _add_paragraph_field(doc, "1. ข้อวิพากษ์ที่สำคัญจากผลการประเมิน โดยนักศึกษาหรือผู้เรียน",
                         _g(data, "eval_crit"))
    _add_paragraph_field(doc, "2. ข้อคิดเห็นของผู้สอนต่อข้อวิพากษ์ตาม ข้อ 1",
                         _g(data, "teacher_comment"))
    _add_paragraph_field(doc, "3. แผนการปรับปรุงการจัดการเรียนรู้ของรายวิชา",
                         _g(data, "improve_plan", "improvement_plan"))
    _add_paragraph_field(doc, "4. ข้อเสนอแนะของผู้สอนต่อคณะกรรมการบริหารหลักสูตร",
                         _g(data, "suggest_to_committee"))

    sigs = ctx.get("signatures", {})
    inst_name = _g(data, "instructors", default=ctx.get("instructor", ""))
    _add_signatures(doc, [
        ("อาจารย์ผู้สอน", inst_name,
         sigs.get(str(inst_name).strip()) or ctx.get("instructor_signature", "")),
    ])

    return _finalize(doc)
