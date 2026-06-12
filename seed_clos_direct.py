"""Seed CourseCLOs for เทคโนโลยีมัลติมีเดีย 2565 using google.cloud.firestore directly.

Requires authorized_user ADC at:
~/.config/gcloud/legacy_credentials/wanchatpookhuntod@gmail.com/adc.json
"""
import json
import os
import re
import sys
import uuid
from collections import defaultdict
from datetime import datetime, timezone

from docx import Document
from google.cloud import firestore as gcp_firestore
from google.oauth2.credentials import Credentials

ADC_PATH = os.path.expanduser(
    "~/.config/gcloud/legacy_credentials/wanchatpookhuntod@gmail.com/adc.json"
)
PROJECT_ID = "qualificationsframework"
PROGRAM_KEYWORD = "มัลติมีเดีย"
PROGRAM_YEAR = 2565
PLO_HEADERS = ["PLO1", "PLO2", "PLO3", "PLO4", "PLO5", "PLO6", "PLO7", "PLO8"]

DOCX_MAIN = "CLO/clo_วท_เทคโนโลยีมัลติมีเดีย.docx"
DOCX_EXTRA = "CLO/clo_วท_เทคโนโลยีมัลติมีเดีย_เพิ่มเติม.docx"


# ── PLO inference rules for the extra file ────────────────────────────────────
def _infer_plos(desc: str) -> str:
    d = desc.lower()
    codes = set()
    # PLO1: ความรู้/ทฤษฎี/หลักการ
    if any(k in d for k in ["ความรู้", "ทฤษฎี", "หลักการ", "แนวคิด", "ศึกษา"]):
        codes.add("PLO1")
    # PLO2: เทคนิค/ทักษะ/ปฏิบัติ
    if any(k in d for k in ["ออกแบบ", "สร้าง", "พัฒนา", "ผลิต", "เทคนิค", "ปฏิบัติ", "ทักษะ"]):
        codes.add("PLO2")
    # PLO3: การวิเคราะห์/แก้ปัญหา
    if any(k in d for k in ["วิเคราะห์", "แก้ปัญหา", "ประเมิน", "ตัดสิน"]):
        codes.add("PLO3")
    # PLO4: จริยธรรม/จรรยาบรรณ
    if any(k in d for k in ["จริยธรรม", "จรรยาบรรณ", "ความรับผิดชอบ", "คุณธรรม"]):
        codes.add("PLO4")
    # PLO5: ความสัมพันธ์/การสื่อสาร
    if any(k in d for k in ["สื่อสาร", "นำเสนอ", "ทำงานร่วม", "ทีม", "ประสานงาน"]):
        codes.add("PLO5")
    # PLO6: เทคโนโลยี/เครื่องมือ
    if any(k in d for k in ["โปรแกรม", "ซอฟต์แวร์", "เครื่องมือ", "เทคโนโลยี", "ระบบ"]):
        codes.add("PLO6")
    # PLO7: ความคิดสร้างสรรค์
    if any(k in d for k in ["สร้างสรรค์", "นวัตกรรม", "ความคิด", "ประยุกต์"]):
        codes.add("PLO7")
    # PLO8: วิชาชีพ/อาชีพ
    if any(k in d for k in ["วิชาชีพ", "อาชีพ", "มาตรฐาน", "อุตสาหกรรม"]):
        codes.add("PLO8")
    if not codes:
        codes.add("PLO1")  # fallback
    return ",".join(sorted(codes))


# ── Docx parsers ──────────────────────────────────────────────────────────────
def parse_main_docx(path):
    """Parse primary CLO docx (has Wingdings PLO checkmarks)."""
    doc = Document(path)
    table = doc.tables[0]
    rows_out = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        clo_col = cells[2] if len(cells) > 2 else ""
        if not clo_col.startswith("CLO"):
            continue
        code = cells[0].strip()
        desc = cells[3].strip() if len(cells) > 3 else ""
        m = re.search(r"\d+", clo_col)
        number = int(m.group()) if m else None
        plos = []
        for idx, cell in enumerate(row.cells[4:], 0):
            if idx >= len(PLO_HEADERS):
                break
            if "<w:sym" in cell._tc.xml:
                plos.append(PLO_HEADERS[idx])
        rows_out.append((code, number, desc, ",".join(plos)))
    return rows_out


def parse_extra_docx(path):
    """Parse extra CLO docx (no Wingdings; infer PLOs from description)."""
    doc = Document(path)
    table = doc.tables[0]
    rows_out = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        clo_col = cells[2] if len(cells) > 2 else ""
        if not clo_col.startswith("CLO"):
            continue
        code = cells[0].strip()
        desc = cells[3].strip() if len(cells) > 3 else ""
        m = re.search(r"\d+", clo_col)
        number = int(m.group()) if m else None
        rows_out.append((code, number, desc, _infer_plos(desc)))
    return rows_out


# ── Firestore helpers ─────────────────────────────────────────────────────────
def get_client():
    with open(ADC_PATH) as f:
        info = json.load(f)
    creds = Credentials(
        token=None,
        refresh_token=info["refresh_token"],
        token_uri="https://oauth2.googleapis.com/token",
        client_id=info["client_id"],
        client_secret=info["client_secret"],
    )
    return gcp_firestore.Client(project=PROJECT_ID, credentials=creds)


def normalize_code(code: str) -> str:
    return code.replace(" ", "").upper()


def seed_clos(db, all_rows):
    # Find program
    programs_ref = db.collection("programs")
    programs = [d for d in programs_ref.stream()]
    target = None
    for doc in programs:
        data = doc.to_dict()
        if PROGRAM_KEYWORD in (data.get("name") or "") and data.get("year") == PROGRAM_YEAR:
            target = (doc.id, data)
            break
    if not target:
        print("ERROR: ไม่พบหลักสูตรเทคโนโลยีมัลติมีเดีย 2565")
        return

    program_id, program_data = target
    print(f"Program: [{program_id}] {program_data.get('name')} {program_data.get('year')}")

    # Fetch courses
    courses_ref = db.collection("courses").where("program_id", "==", program_id)
    courses = list(courses_ref.stream())
    course_map = {}
    for doc in courses:
        data = doc.to_dict()
        code = normalize_code(data.get("code") or "")
        if code:
            course_map[code] = (doc.id, data)
    print(f"Loaded {len(courses)} courses from program")

    # Group rows by course code
    by_course = defaultdict(list)
    for (code, number, desc, plos) in all_rows:
        by_course[normalize_code(code)].append((number, desc, plos))

    inserted = 0
    skipped = []
    clo_col = db.collection("course_clos")

    for norm_code, clo_list in by_course.items():
        if norm_code not in course_map:
            skipped.append(norm_code)
            continue
        course_id = course_map[norm_code][0]

        # Delete existing CLOs
        existing = list(clo_col.where("course_id", "==", course_id).stream())
        for doc in existing:
            doc.reference.delete()

        for (number, desc, plos) in clo_list:
            clo_col.document(str(uuid.uuid4())).set({
                "course_id": course_id,
                "code": f"CLO{number}" if number else "CLO?",
                "description": desc,
                "plo_codes": plos,
                "order": number,
                "created_at": datetime.now(timezone.utc),
            })
            inserted += 1

    print(f"\nInserted {inserted} CLOs")
    if skipped:
        print(f"Skipped (course not found): {skipped}")


def main():
    rows = []
    if os.path.exists(DOCX_MAIN):
        main_rows = parse_main_docx(DOCX_MAIN)
        print(f"Parsed {len(main_rows)} CLO rows from main docx")
        rows.extend(main_rows)
    else:
        print(f"WARNING: {DOCX_MAIN} not found")

    if os.path.exists(DOCX_EXTRA):
        extra_rows = parse_extra_docx(DOCX_EXTRA)
        print(f"Parsed {len(extra_rows)} CLO rows from extra docx")
        rows.extend(extra_rows)
    else:
        print(f"WARNING: {DOCX_EXTRA} not found")

    if not rows:
        print("No rows parsed – exiting")
        return

    db = get_client()
    seed_clos(db, rows)


if __name__ == "__main__":
    main()
