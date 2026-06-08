"""Smoke tests for the Word (.docx) exporters.

Pure-Python: no Firestore access required. They confirm the builders produce a
valid, openable docx and that Thai text plus indexed/array fields land in it.
"""

import io

from docx import Document

from exporters import build_tqf3_docx, build_tqf4_docx, build_tqf5_docx

CTX = {
    "faculty": "วิทยาศาสตร์",
    "program": "มัลติมีเดีย",
    "course_code": "CS101",
    "course_name": "การเขียนโปรแกรม",
    "credits": 3,
    "description": "พื้นฐานการเขียนโปรแกรม",
    "section_number": "1",
    "semester": 1,
    "year": 2567,
    "instructor": "อ.วันชาติ",
}


def _all_text(buffer: io.BytesIO) -> str:
    doc = Document(io.BytesIO(buffer.getvalue()))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    # Strip the U+200B Thai word-break hints so content assertions stay readable.
    return "\n".join(parts).replace("​", "")


def test_build_tqf3_docx_includes_array_fields():
    general = {
        "course_objective": "เข้าใจพื้นฐาน",
        "clo_text[]": ["CLO1 เขียนโปรแกรมได้", "CLO2 แก้ปัญหาได้"],
        "plo[]": ["PLO1"],
        "teach_strategy[]": ["บรรยาย"],
        "assess_strategy[]": ["สอบ"],
        "topic[]": ["บทนำ", "ตัวแปร"],
        "references": "APA references",
    }
    text = _all_text(build_tqf3_docx(general, CTX))
    assert "มคอ.3" in text
    assert "CS101" in text  # ctx fallback for course_code
    assert "CLO1 เขียนโปรแกรมได้" in text
    assert "บทนำ" in text


def test_build_tqf3_docx_falls_back_to_context():
    # Empty payload should still render with course/term context.
    text = _all_text(build_tqf3_docx({}, CTX))
    assert "วิทยาศาสตร์" in text  # faculty fallback
    assert "การเขียนโปรแกรม" in text  # course_name fallback


def test_build_tqf5_docx_includes_indexed_clo_and_grades():
    data = {
        "clo_desc_1": "CLO1",
        "clo_teach_1": "บรรยาย",
        "clo_desc_2": "CLO2",
        "n_registered": 40,
        "g_level_1": "A",
        "g_count_1": 10,
        "issue_1": "เวลาน้อย",
        "issue_fix_1": "เพิ่มคาบ",
    }
    text = _all_text(build_tqf5_docx(data, CTX))
    assert "รายงานผลการจัดการเรียนรู้ระดับรายวิชา" in text
    assert "CLO1" in text
    assert "CLO2" in text
    assert "เวลาน้อย" in text


def test_build_tqf5_docx_matches_new_format_title_major_and_order():
    data = {
        "major": "เทคโนโลยีมัลติมีเดีย",
        "fail_fix": "จัดสอนเสริม",
        "verification": "ทวนสอบโดยกรรมการหลักสูตร",
    }
    text = _all_text(build_tqf5_docx(data, CTX))
    # New format title includes "ของหลักสูตร".
    assert "รายงานผลการจัดการเรียนรู้ระดับรายวิชาของหลักสูตร" in text
    # Major (วิชาเอก) is now rendered for TQF5 too.
    assert "เทคโนโลยีมัลติมีเดีย" in text
    # Verification (format item 8) comes after the fail-fix block, not before grades.
    assert text.index("ทวนสอบโดยกรรมการหลักสูตร") > text.index("จัดสอนเสริม")


def test_build_tqf3_docx_includes_rubric_and_new_columns():
    general = {
        "course_status": "วิชาบังคับ",
        "plos": "PLO1 PLO2",
        "clo_text[]": ["CLO1"],
        "teach_strategy[]": ["บรรยาย"],
        "assess_criteria[]": ["ผ่านเกณฑ์ร้อยละ 60"],
        "assess_strategy[]": ["สอบ"],
        "rubric_topic[]": ["ทักษะการสื่อสาร"],
        "rubric_l5[]": ["ดีมาก"],
        "appeal_channel": "ติดต่ออาจารย์ทาง email",
    }
    text = _all_text(build_tqf3_docx(general, CTX))
    assert "เกณฑ์การวัดและการประเมินผล" in text
    assert "ผ่านเกณฑ์ร้อยละ 60" in text
    assert "Rubric Score" in text
    assert "ทักษะการสื่อสาร" in text


def test_build_tqf4_docx_field_experience():
    general = {
        "field_objective": "ฝึกปฏิบัติงานจริง",
        "report[]": ["รายงานสรุปการฝึก"],
        "report_criteria[]": ["ผ่าน"],
        "clo_text[]": ["CLO1 ปฏิบัติงานได้"],
        "assess_clo[]": ["CLO1"],
        "assess_evaluator[]": ["พี่เลี้ยง"],
        "mentor_duties": "ดูแลนักศึกษา",
    }
    text = _all_text(build_tqf4_docx(general, CTX))
    assert "มคอ.4" in text
    assert "ประสบการณ์ภาคสนาม" in text
    assert "CLO1 ปฏิบัติงานได้" in text
    assert "พี่เลี้ยง" in text


def test_exporters_strip_xml_illegal_control_chars():
    # Users paste from PDFs/Word, smuggling control chars (\x00, \x0b, \x0c) into
    # the form. python-docx raises ValueError on those, 500-ing the export. Every
    # builder must strip them while keeping legal whitespace (tab/newline).
    nasty = "นักศึกษา\x0bสามารถ\x00อธิบาย\x0cได้\x07\x1f"
    # The course header (code/name) flows through every builder, so the ctx
    # course_name is a reliable place to prove illegal chars are scrubbed.
    ctx = {**CTX, "course_name": nasty}
    for build in (build_tqf3_docx, build_tqf4_docx, build_tqf5_docx):
        text = _all_text(build({}, ctx))
        assert "นักศึกษาสามารถอธิบายได้" in text
        for ctrl in ("\x00", "\x0b", "\x0c", "\x07", "\x1f"):
            assert ctrl not in text


def test_is_field_experience_course_detection():
    from types import SimpleNamespace

    from app import is_field_experience_course

    field = SimpleNamespace(name_th="การฝึกประสบการณ์วิชาชีพครู", name_en="", code="ED401")
    coop = SimpleNamespace(name_th="", name_en="Cooperative Education", code="CS490")
    normal = SimpleNamespace(name_th="การเขียนโปรแกรม", name_en="Programming", code="CS101")

    assert is_field_experience_course(field) is True
    assert is_field_experience_course(coop) is True
    assert is_field_experience_course(normal) is False
    assert is_field_experience_course(None) is False
